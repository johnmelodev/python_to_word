# pip3 install python-docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Cria um novo Documento Word
doc = Document()

# Define o nome e o alinhamento do título
font_size_title = Pt(11)  # Define o tamanho da fonte para 11
title = doc.add_heading(level=1)
run = title.add_run('Currículo de João Melo')
run.bold = True
run.font.name = 'Nunito'
run.font.size = font_size_title
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Adiciona informações de contato
font_size_paragraph = Pt(9)  # Define o tamanho da fonte para 9 para os parágrafos
doc.add_heading('Contato', level=2)
paragraph = doc.add_paragraph(style='BodyText')
run = paragraph.add_run('LinkedIn | 353899604104 | My Website | joaohenrique1231500@gmail.com | GitHub')
run.font.name = 'Nunito'
run.font.size = font_size_paragraph

# Adiciona habilidades
doc.add_heading('Skills', level=2)
doc.add_paragraph('• JavaScript | Node.js | React.js | Nest.js | jQuery | Python | Django | Flask| SQLite | SQLAlchemy | Oracle | MySQL | Git\n• Web Development | Frontend | Backend | Full-Stack | HTML | CSS | Bootstrap | Sitecore | Wix\n• Others: Crawlers, Chatbots, Selenium, AI Prompt Creation| Languages: Portuguese (Native), English (Fluent/Native), Spanish (Basic)', style='BodyText')

# Adiciona objetivos
doc.add_heading('Objetive', level=2)
doc.add_paragraph('• To work as a Front-end Engineer (ReactJS, NestJS, TypeScript), and in web systems. Contribute to a significant project, and share ideas with a smart team..', style='BodyText')

# Adiciona experiência de trabalho
doc.add_heading('Experience', level=2)
doc.add_paragraph('• Front-end Developer at Valtech Inc., Remote, Worldwide, 03/2021 - 03/2024\n• Implemented innovative web solutions for a L’Oréal Paris campaign, playing a key role in boosting Maybelline sales by 35% and leading a successful quiz campaign, featuring actress Drew Barrymore.\n• Utilizing my React.js skills, I crafted components for Sitecore CMS, standardizing operations across various brands including Mandarin Orientals Hotels and L’Oréal Paris.\n• I supported the integration of backend functionality and data using Node.js into the frontend, closely collaborating with backend teams.\n• Participated in an urgent project, where my team and I rebuilt the global websites of Mandarin Oriental. With the integration of Sitecore and React.js technologies, deploying the year-long project within the three-month deadline set by the client.(38 hotels in 25 countries).\n• Became a Senior Sitecore Developer and received pay raises every year due to the Valtech ⭐ Award for excellent feedback.\n• Software Engineer/ Django Developer at Companhia do Papel, Londrina, PR, Brazil, 02/2019 - 02/2021\n• Implemented a login and logout system for their website using Django and PostgreSQL, with data stored in Supabase. Front-end made with Bootstrap.\n• In addition, I developed a software running on VPS (Linode), to scrapy data from a list of supplier websites using Selenium, and transferring it to Google Sheets equipped with formulas for profit calculation via Google Sheets API.\n• Data scraping automation has led to a huge increase in inventory accuracy and profit, saving around 15 hours of manual work per week.\n• My entry into the tech industry was triggered by a friend’s recommendation, who was impressed by my Selenium projects. After it I worked in that big store complex with a partnership with Base 2 tech company.\n• Since 2019, I’ve continuously provided services across the nation as a Software Engineer, Django Developer, and Web Designer.', style='BodyText')

# Adiciona projetos técnicos
doc.add_heading('Technical Projects', level=2)
doc.add_paragraph('• Django Investing list: Developed a full-stack application for managing investment data. Utilized Django, PostgreSQL database, and implemented a secure login system and a user-friendly interface with Bootstrap.\n• Price Scrapy to Excel to WhatsApp: Python-based automation for price scraping (Selenium) and profit calculation in Excel (Openpyxl). Delivers key insights on profit margins and sends data to WhatsApp using Pyautogui.\n• Django Price-Search: Developed a Django-based system for competitive price analysis, Selenium scraping various websites and presenting results on a hosted platform.\n• Telegram Bot for Financial Data Management: Developed a Python-based Telegram bot integrated with Google Sheets API for automated financial management. Simplifies tracking of income, expenses, and transfers.', style='BodyText')

# Adiciona educação
doc.add_heading('Education and Courses', level=2)
doc.add_paragraph('Bachelor of Administration, UEL - State University of Londrina, Londrina, PR, Brazil, 03/2016 - 12/2021', style='BodyText')

# Adiciona mentorship
doc.add_heading('Mentorship', level=2)
doc.add_paragraph('• DevAprender & DoDev Schools: Volunteered as a mentor in Django, Python automation, Fullstack, and JavaScript courses, providing code troubleshooting and helping with technical doubts in the platform.', style='BodyText')

# Adiciona others
doc.add_heading('Others', level=2)
doc.add_paragraph('• Proficient in team project technologies like Miro, Jira Boards, simultaneously excelling in over 5 projects within the same consultancy company.\n• Completed high school in Portugal; over a year of professional experience in the United States under J1 status.', style='BodyText')

# Salva o documento
doc.save('curriculo.docx')
